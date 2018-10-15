#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import argparse
from docx import Document

from section import SectionList
from class_id import ClassIdList
from text import camelcase


MEClassSection = "11.2.4"       # Class IDs


def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Final Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='ITU-T G.988-201711.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.PreCompiled.json',
                        help='Path to pre-parsed G.988 data, default: G.988.PreCompiled.json')

    parser.add_argument('--output', '-o', action='store',
                        default='G.988.Parsed.json',
                        help='Output filename, default: G.988.Parsed.json')

    parser.add_argument('--classes', '-c', action='store',
                        default='11.2.4',
                        help='Document section number with ME Class IDs, default: 11.2.4')

    args = parser.parse_args()
    return args


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        self.sections = None
        self.class_ids = None
        self.body = None

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        document = self.load_itu_document()
        self.sections = SectionList()
        self.sections.load(self.args.input)

        self.paragraphs = document.paragraphs
        # doc_sections = document.sections
        # styles = document.styles
        # self.body = document.element.body

        print('Extracting ME Class ID values')
        self.class_ids = ClassIdList.parse_sections(self.sections,
                                                    self.args.classes)

        print('Found {} ME Class ID entries. {} have sections associated to them'.
              format(len(self.class_ids),
                     len([c for c in self.class_ids.values()
                          if c.section is not None])))

        crazy_formatted_mes = \
            {23,
             157,           # Large String
             }
        print('Skipping the following MEs due to complex document formatting')
        print("    {}".format(crazy_formatted_mes))
        self.class_ids = {k: v for k, v in self.class_ids.items()
                          if k not in crazy_formatted_mes}

        print('Managed Entities without Sections')
        for c in [c for c in self.class_ids.values() if c.section is None]:
            print('    {:>4}: {}'.format(c.cid, c.name))

        # Work with what we have
        self.class_ids = {cid: c for cid, c in self.class_ids.items()
                          if c.section is not None}
        print('')
        print('Parsing deeper for managed Entities with Sections')
        for c in self.class_ids.values():
            print('    {:>9}:  {:>4}: {} -> {}'.format(c.section.section_number,
                                                       c.cid,
                                                       c.name,
                                                       camelcase(c.name)))
            c.deep_parse(self.paragraphs)

        completed = len([c.state == 'complete' for c in self.class_ids.values()])
        failed = len([c.state == 'failure' for c in self.class_ids.values()])

        print('Of {} MEs, {} were parsed successfully and {} failed'.format(len(self.class_ids),
                                                                            completed,
                                                                            failed))
        # Run some sanity checks
        print('\n\n\nValidating ME Class Information:\n')
        for c in self.class_ids.values():
            print('  Class ID: {} - {}'.format(c.cid, c.name))
            if len(c.attributes == 0):
                print('    NO ATTRIBUTES')

            for attr in c.attributes:
                if attr.access is None or len(attr.access) == 0:
                    print('    NO ACCESS INFORMATION')
                # if attr.size is None:
                #     print('    NO SIZE INFORMATION')      TODO: Get Size decode working

        # Output the results to a JSON file so it can be used by a code-generation
        # tool

        # TODO: Write it out here.


if __name__ == '__main__':
    Main().start()
