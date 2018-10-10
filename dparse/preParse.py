#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

#
#   This program will walk the ITU G.988 docx document and extract section
#   headers (and following paragraphs) and any Table Items.  It then
#   will save this to a pre-compiled JSON file that can then be used
#   later to quickly load the sections for further processing

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import sys
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from section import SectionHeading, SectionList
from tables import TableList, Table


def parse_args():
    return {
        'input': 'T-REC-G.988-201711-I!!MSW-E.docx',
        'output': 'G.988.PreCompiled.json',
    }


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()

    def start(self):
        sections = SectionList()
        document = Document(self.args['input'])

        paragraphs = document.paragraphs
        doc_sections = document.sections
        styles = document.styles
        tables = document.tables

        print('Number of sections  : {}'.format(len(doc_sections)))
        print('Number of paragraphs: {}'.format(len(paragraphs)))
        print('Number of styles    : {}, {} are built-in styles'.format(len(styles),
                                                                        len([x for x in styles if x.builtin])))
        print('Number of tables    : {}'.format(len(tables)))
        print('Parsing paragraphs & tables to extract high level info. This will take three+ minutes')

        pnum = 0
        tnum = 0
        current_section = None

        def is_section_header(p):
            return (isinstance(p, Paragraph)
                    and len(p.text)
                    and p.style.builtin
                    and 'heading ' in p.style.name.lower())

        for block in Main.iter_block_items(document):
            if isinstance(block, Paragraph):
                if is_section_header(block):
                    # Save of previous
                    current_section = SectionHeading.create(pnum, block)
                    sections.add(current_section)

                elif len(block.text) > 0 and current_section is not None:
                    current_section.add_contents(pnum)

                pnum += 1

            elif isinstance(block, DocxTable):
                if current_section is not None:
                    table = Table.create(tnum, block)
                    current_section.add_contents(table)
                tnum += 1

            else:
                print('Unsupported block type: {}', type(block))

            if pnum % 25 == 24:
                print('.', end='')
                sys.stdout.flush()
            if pnum % 2000 == 1999:
                print('')

        # Save to file
        print('')
        print('Saving Section parsing information to {}'.format(self.args['output']))
        sections.save(self.args['output'])

        print('Section pre-parsing are complete')

        # Restore and verify
        sections.load(self.args['output'])

        for section in sections:
            print('  Section: {} -> {}'.format(section, section.section_points))

    @staticmethod
    def iter_block_items(parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
            # print(parent_elm.xml)
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield DocxTable(child, parent)


if __name__ == '__main__':
    try:
        Main().start()

    except Exception as _e:
        raise
